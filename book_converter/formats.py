"""
Format-specific converters for transforming PDF content to various file formats.
"""
import os
import logging
from abc import ABC, abstractmethod
from typing import Dict, Any, List, Optional, Tuple, TYPE_CHECKING
from pathlib import Path
import tempfile
import shutil
import re
import subprocess

from .chapter_patterns import ChapterDetector, CHAPTER_PATTERNS, CHAPTER_STYLES

# This prevents circular imports
if TYPE_CHECKING:
    from .converter import PDFConverter

# Third-party libraries
import ebooklib
from ebooklib import epub
import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import markdown
from PIL import Image

# Configure logger
logger = logging.getLogger(__name__)


class BaseFormatConverter(ABC):
    """
    Base abstract class for all format converters.
    """
    
    def __init__(self, pdf_converter: 'PDFConverter'):
        """
        Initialize the format converter with a PDFConverter instance.
        
        Args:
            pdf_converter (PDFConverter): PDFConverter instance
        """
        self.pdf_converter = pdf_converter
    
    @abstractmethod
    def convert(self, output_path: str, **kwargs) -> str:
        """
        Convert the PDF to the target format.
        
        Args:
            output_path (str): Path to save the output file
            **kwargs: Format-specific options
            
        Returns:
            str: Path to the output file
        """
        pass
        

class EPUBConverter(BaseFormatConverter):
    """
    Converter for EPUB format.
    """
    
    def convert(self, output_path: str, **kwargs) -> str:
        """
        Convert the PDF to EPUB format.
        
        Args:
            output_path (str): Path to save the output EPUB file
            **kwargs: Additional options:
                - title (str): Book title
                - author (str): Book author
                - language (str): Book language
                - cover_image (str): Path to cover image
                - css (str): Custom CSS
                - toc_depth (int): Table of contents depth
                - strip_text (list): Text patterns to strip
                - chapter_pattern (str): Regex pattern to detect chapter openings
                - chapter_style (str): CSS style for chapter openings
                - chapter_style_name (str): Predefined style name ('standard', 'quoted', 'decorative')
                
        Returns:
            str: Path to the output EPUB file
        """
        try:
            strip_text = kwargs.get('strip_text') or []
            
            # Extract text from the PDF
            text_by_page_ = self.pdf_converter.extract_text(
                include_tables=kwargs.get('include_tables', True),
                detect_columns=kwargs.get('detect_columns', True)
            )
            
            text_by_page = []
            for page in text_by_page_:
                # Remove unwanted characters
                page = re.sub(r'\s+', ' ', page)
                page = re.sub(r'^\s+|\s+$', '', page)
                page = re.sub(r'\u200b', '\n', page)
                
                if strip_text:
                    for pattern in strip_text:
                        page = re.sub(pattern, '', page)
                        
                # Add cleaned page text to the list
                text_by_page.append(page)
                        
            # Create a new EPUB book
            book = epub.EpubBook()
            
            # Handle chapter pattern and styling
            chapter_pattern = kwargs.get('chapter_pattern')
            chapter_style_name = kwargs.get('chapter_style_name')
            chapter_style = kwargs.get('chapter_style')
            
            # If a predefined style name is provided, use that
            if chapter_style_name and chapter_style_name in CHAPTER_STYLES:
                chapter_style = CHAPTER_STYLES[chapter_style_name]
            
            # If a predefined pattern name is provided, use that
            if isinstance(chapter_pattern, str) and chapter_pattern in CHAPTER_PATTERNS:
                chapter_pattern = CHAPTER_PATTERNS[chapter_pattern]
            
            # Initialize chapter detector if pattern is provided
            chapter_detector = None
            if chapter_pattern:
                try:
                    chapter_detector = ChapterDetector(chapter_pattern)
                    logger.info(f"Chapter pattern initialized: {chapter_pattern}")
                except Exception as e:
                    logger.warning(f"Failed to initialize chapter detector: {e}")
            
            # Set metadata
            title = kwargs.get('title') or self.pdf_converter.metadata.get('title', 'Untitled Book')
            author = kwargs.get('author') or self.pdf_converter.metadata.get('author', 'Unknown Author')
            language = kwargs.get('language', 'en')
            
            book.set_title(title)
            book.set_language(language)
            book.add_author(author)
            
            # Add cover if provided
            cover_image_path = kwargs.get('cover_image')
            if cover_image_path and os.path.exists(cover_image_path):
                with open(cover_image_path, 'rb') as img_file:
                    cover_content = img_file.read()
                    book.set_cover('cover.jpg', cover_content)
            elif self.pdf_converter.images and len(self.pdf_converter.images) > 0:
                # Use the first image as cover if available
                first_image = self.pdf_converter.images[0]
                if 'image' in first_image:
                    book.set_cover('cover.jpg', first_image['image'])
                
            # Create CSS
            default_css = '''
            /* Chapter opening styles */
            .chapter-opening {
                margin: 2em 0;
                font-style: italic;
                text-align: center;
                line-height: 1.6;
            }
            
            /* Base Typography */
            body {
                font-family: 'Palatino', 'Palatino Linotype', 'Book Antiqua', serif;
                font-size: 1em;
                line-height: 1.6;
                margin: 5% 6%;
                text-align: justify;
                color: #333;
                max-width: 40em;
                hyphens: auto;
            }

            /* Headings */
            h1, h2, h3, h4, h5, h6 {
                font-family: 'Helvetica Neue', 'Arial', sans-serif;
                line-height: 1.2;
                margin-top: 2.5em;
                margin-bottom: 1em;
                color: #222;
                font-weight: 600;
                page-break-after: avoid;
            }
            h1 {
                font-size: 1.6em;
                text-align: center;
                margin-top: 3em;
                border-bottom: 1px solid #ddd;
                padding-bottom: 0.5em;
            }
            h2 { font-size: 1.4em; }
            h3 { font-size: 1.2em; }
            h4, h5, h6 { font-size: 1em; }

            /* Paragraphs */
            p {
                margin-top: 0.75em;
                margin-bottom: 0.75em;
                orphans: 2;
                widows: 2;
            }

            /* Links */
            a {
                color: #0066cc;
                text-decoration: none;
            }

            /* Lists */
            ul, ol {
                margin: 1em 0 1em 1em;
                padding-left: 1em;
            }
            li {
                margin-bottom: 0.5em;
            }

            /* Images */
            img {
                max-width: 100%;
                height: auto;
                display: block;
                margin: 1.5em auto;
                page-break-inside: avoid;
            }
            figure {
                margin: 1.5em 0;
                text-align: center;
            }
            figcaption {
                font-size: 0.9em;
                font-style: italic;
                margin-top: 0.5em;
                color: #555;
            }

            /* Blockquotes */
            blockquote {
                margin: 1.5em 2em;
                padding-left: 1em;
                border-left: 3px solid #ddd;
                font-style: italic;
                color: #555;
            }

            /* Tables */
            table {
                width: 100%;
                margin: 1.5em 0;
                border-collapse: collapse;
                font-size: 0.9em;
            }
            th, td {
                padding: 0.5em;
                border: 1px solid #ddd;
            }
            th {
                background-color: #f8f8f8;
                font-weight: bold;
            }

            /* Code and Pre */
            code, pre {
                font-family: "Courier New", Courier, monospace;
                font-size: 0.9em;
                white-space: pre-wrap;
                background-color: #f8f8f8;
                padding: 0.2em 0.4em;
                border-radius: 3px;
            }
            pre {
                margin: 1em 0;
                padding: 1em;
                overflow-x: auto;
                border: 1px solid #ddd;
                line-height: 1.4;
            }
            pre code {
                padding: 0;
                background: none;
            }

            /* Horizontal Rule */
            hr {
                height: 1px;
                border: 0;
                background-color: #ddd;
                margin: 2em 0;
            }

            /* Special Elements */
            .footnote {
                font-size: 0.9em;
                color: #666;
                vertical-align: super;
                line-height: 0;
            }
            .caption {
                font-size: 0.9em;
                font-style: italic;
                text-align: center;
                color: #555;
            }
            .page-break {
                page-break-before: always;
            }
            '''
            # Add chapter style to CSS if provided
            if chapter_style:
                default_css += chapter_style
                
            css = kwargs.get('css', default_css)
            css_file = epub.EpubItem(
                uid="style_default",
                file_name="style/default.css",
                media_type="text/css",
                content=css
            )
            book.add_item(css_file)
            
            # Create chapters
            chapters = []
            toc = []
            
            # Process table of contents if available
            pdf_toc = self.pdf_converter.toc
            has_toc = pdf_toc and len(pdf_toc) > 0
            
            # If we have TOC, use it to structure chapters
            if has_toc:
                current_page = 0
                for i, (level, title, page) in enumerate(pdf_toc):
                    if title in (' ', '- ', '\u200b '):
                        continue
                    
                    next_page = pdf_toc[i + 1][2] if i + 1 < len(pdf_toc) else len(text_by_page)
                    
                    # Get content for this chapter
                    chapter_content = "".join(text_by_page[page-1:next_page-1])
                    
                    # Apply chapter pattern formatting if detector is available
                    if chapter_detector:
                        chapter_content = chapter_detector.format_chapter_openings(chapter_content)
                    
                    # Create chapter
                    chapter = self._create_chapter(f'chapter_{i}', title, chapter_content, css_file)
                    book.add_item(chapter)
                    chapters.append(chapter)
                    
                    # Add to toc
                    if level <= kwargs.get('toc_depth', 3):
                        toc.append((chapter, chapter.file_name))
            else:
                # No TOC, create a chapter per page or group pages
                page_count = len(text_by_page)
                
                # Group pages into chapters if there are too many
                if page_count > 30:
                    chunk_size = max(10, page_count // 20)  # Aim for about 20 chapters
                    for i in range(0, page_count, chunk_size):
                        chapter_num = i // chunk_size + 1
                        end_idx = min(i + chunk_size, page_count)
                        
                        chapter_content = "".join(text_by_page[i:end_idx])
                                
                        chapter = self._create_chapter(
                            f'chapter_{chapter_num}',
                            f'Chapter {chapter_num}',
                            chapter_content,
                            css_file
                        )
                        book.add_item(chapter)
                        chapters.append(chapter)
                        toc.append((chapter, chapter.file_name))
                else:
                    # Create a chapter for each page if there aren't too many
                    for i, page_text in enumerate(text_by_page):
                        chapter = self._create_chapter(
                            f'page_{i+1}',
                            f'Page {i+1}',
                            page_text,
                            css_file
                        )
                        book.add_item(chapter)
                        chapters.append(chapter)
                        toc.append((chapter, chapter.file_name))
            
            # Add images
            for i, img_data in enumerate(self.pdf_converter.images):
                if 'image' in img_data:
                    img_item = epub.EpubItem(
                        uid=f'image_{i}',
                        file_name=f'images/image_{i}.{img_data["ext"]}',
                        media_type=f'image/{img_data["ext"]}',
                        content=img_data['image']
                    )
                    book.add_item(img_item)
            
            # Define Table of Contents
            try:
                book.toc = toc
                logger.info(f"Added {len(toc)} items to table of contents")
            except Exception as toc_error:
                logger.warning(f"Error setting TOC: {toc_error}. Creating simplified TOC.")
                # Create a simplified TOC as fallback
                simplified_toc = []
                for chapter in chapters:
                    try:
                        simplified_toc.append((chapter, chapter.file_name))
                    except Exception:
                        pass
                book.toc = simplified_toc
            
            # Add default NCX and Nav
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())
            
            # Define spine
            book.spine = ['nav'] + chapters
            
            # Write to file
            epub.write_epub(output_path, book, {})
            
            logger.info(f"Successfully converted to EPUB: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error converting to EPUB: {e}")
            raise ValueError(f"Failed to convert to EPUB: {e}")
    
    def _create_chapter(self, id_name, title, content, css):
        """
        Create an EPUB chapter.
        
        Args:
            id_name (str): Unique ID for the chapter
            title (str): Chapter title
            content (str): Chapter content
            css (epub.EpubItem): CSS file
            
        Returns:
            epub.EpubHtml: Chapter object
        """
        chapter = epub.EpubHtml(
            title=title, 
            file_name=f'{id_name}.xhtml',
            lang='en'
        )
        chapter.content = f'''
        <html>
        <head>
            <title>{title}</title>
            <link rel="stylesheet" href="style/default.css" type="text/css" />
        </head>
        <body>
            <h1>{title}</h1>
            {content}
        </body>
        </html>
        '''
        chapter.add_item(css)
        return chapter


class DOCXConverter(BaseFormatConverter):
    """
    Converter for DOCX format.
    """
    
    def convert(self, output_path: str, **kwargs) -> str:
        """
        Convert the PDF to DOCX format.
        
        Args:
            output_path (str): Path to save the output DOCX file
            **kwargs: Additional options:
                - title (str): Document title
                - include_images (bool): Whether to include images
                - heading_style (str): Style for headings
                
        Returns:
            str: Path to the output DOCX file
        """
        try:
            # Extract text and data from the PDF
            text_by_page = self.pdf_converter.extract_text(
                include_tables=kwargs.get('include_tables', True),
                detect_columns=kwargs.get('detect_columns', True)
            )
            
            # Create a new DOCX document
            doc = Document()
            
            # Set metadata
            title = kwargs.get('title') or self.pdf_converter.metadata.get('title', 'Untitled Document')
            doc.core_properties.title = title
            
            if self.pdf_converter.metadata.get('author'):
                doc.core_properties.author = self.pdf_converter.metadata.get('author')
                
            # Add title page if requested
            if kwargs.get('add_title_page', True):
                title_paragraph = doc.add_paragraph()
                title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = title_paragraph.add_run(title)
                title_run.font.size = Pt(24)
                title_run.font.bold = True
                
                # Add author if available
                if self.pdf_converter.metadata.get('author'):
                    author_paragraph = doc.add_paragraph()
                    author_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    author_paragraph.add_run(self.pdf_converter.metadata.get('author'))
                    
                doc.add_page_break()
            
            # Process table of contents if available
            pdf_toc = self.pdf_converter.toc
            has_toc = pdf_toc and len(pdf_toc) > 0
            
            # Add content based on TOC or page by page
            if has_toc:
                # Use TOC to structure the document
                current_page = 0
                for i, (level, title, page) in enumerate(pdf_toc):
                    # Add heading with appropriate level
                    heading = doc.add_heading(title, level=min(level, 9))
                    
                    # Get content for this section
                    next_page = pdf_toc[i+1][2] if i+1 < len(pdf_toc) else len(text_by_page)
                    section_content = "".join(text_by_page[page-1:next_page-1])
                    
                    # Break into paragraphs and add
                    paragraphs = section_content.split('\n\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            doc.add_paragraph(para_text.strip())
            else:
                # No TOC, just add content page by page
                for i, page_text in enumerate(text_by_page):
                    # Break into paragraphs and add
                    paragraphs = page_text.split('\n\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            doc.add_paragraph(para_text.strip())
                    
                    # Add page break between pages
                    if i < len(text_by_page) - 1:
                        doc.add_page_break()
            
            # Add images if requested
            if kwargs.get('include_images', True) and self.pdf_converter.images:
                doc.add_heading('Images', level=1)
                
                # Create a temp directory for images
                with tempfile.TemporaryDirectory() as temp_dir:
                    for i, img_data in enumerate(self.pdf_converter.images):
                        if 'image' in img_data:
                            img_path = os.path.join(temp_dir, f'image_{i}.{img_data["ext"]}')
                            
                            # Save image to temp file
                            with open(img_path, 'wb') as img_file:
                                img_file.write(img_data['image'])
                            
                            # Add to document
                            try:
                                doc.add_picture(img_path, width=Inches(6))
                                doc.add_paragraph(f"Image {i+1}")
                            except Exception as img_err:
                                logger.warning(f"Could not add image {i}: {img_err}")
            
            # Save the document
            doc.save(output_path)
            
            logger.info(f"Successfully converted to DOCX: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error converting to DOCX: {e}")
            raise ValueError(f"Failed to convert to DOCX: {e}")


class HTMLConverter(BaseFormatConverter):
    """
    Converter for HTML format.
    """
    
    def convert(self, output_path: str, **kwargs) -> str:
        """
        Convert the PDF to HTML format.
        
        Args:
            output_path (str): Path to save the output HTML file
            **kwargs: Additional options:
                - title (str): Document title
                - include_images (bool): Whether to include images
                - css (str): Custom CSS
                - embed_images (bool): Whether to embed images as base64
                
        Returns:
            str: Path to the output HTML file
        """
        try:
            # Extract text and data from the PDF
            text_by_page = self.pdf_converter.extract_text(
                include_tables=kwargs.get('include_tables', True),
                detect_columns=kwargs.get('detect_columns', True)
            )
            
            # Set metadata
            title = kwargs.get('title') or self.pdf_converter.metadata.get('title', 'Untitled Document')
            author = self.pdf_converter.metadata.get('author', 'Unknown Author')
            
            # Default CSS
            default_css = '''
            body {
                font-family: serif;
                line-height: 1.5;
                margin: 2em auto;
                max-width: 800px;
                padding: 0 1em;
            }
            h1, h2, h3, h4, h5, h6 {
                font-family: sans-serif;
                margin-top: 1.5em;
                margin-bottom: 0.5em;
            }
            img {
                max-width: 100%;
                height: auto;
                display: block;
                margin: 1em auto;
            }
            .page-break {
                page-break-after: always;
                margin: 2em 0;
                border-bottom: 1px dashed #ccc;
            }
            '''
            css = kwargs.get('css', default_css)
            
            # Create the base HTML structure
            html = f'''<!DOCTYPE html>
            <html lang="en">
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>{title}</title>
                <style>
                {css}
                </style>
            </head>
            <body>
                <header>
                    <h1>{title}</h1>
                    <p>Author: {author}</p>
                </header>
                <main>
            '''
            
            # Process table of contents if available
            pdf_toc = self.pdf_converter.toc
            has_toc = pdf_toc and len(pdf_toc) > 0
            
            # Create TOC section if TOC exists
            if has_toc:
                html += '<nav id="toc"><h2>Table of Contents</h2><ul>'
                
                for level, title, page in pdf_toc:
                    # Indentation based on level
                    indent = '    ' * (level - 1)
                    html += f'{indent}<li><a href="#page_{page}">{title}</a></li>\n'
                
                html += '</ul></nav>\n'
            
            # Add content based on TOC or page by page
            if has_toc:
                # Use TOC to structure the document
                current_page = 0
                for i, (level, title, page) in enumerate(pdf_toc):
                    # Add section heading
                    html += f'<section id="page_{page}">\n'
                    html += f'<h{min(level+1, 6)}>{title}</h{min(level+1, 6)}>\n'
                    
                    # Get content for this section
                    next_page = pdf_toc[i+1][2] if i+1 < len(pdf_toc) else len(text_by_page)
                    section_content = "".join(text_by_page[page-1:next_page-1])
                    
                    # Convert newlines to <p> tags
                    paragraphs = section_content.split('\n\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            html += f'<p>{para_text.strip()}</p>\n'
                    
                    html += '</section>\n'
            else:
                # No TOC, just add content page by page
                for i, page_text in enumerate(text_by_page):
                    html += f'<section id="page_{i+1}">\n'
                    
                    # Convert newlines to <p> tags
                    paragraphs = page_text.split('\n\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            html += f'<p>{para_text.strip()}</p>\n'
                            
                    html += '</section>\n'
                    
                    # Add page break between pages
                    if i < len(text_by_page) - 1:
                        html += '<div class="page-break"></div>\n'
            
            # Handle images
            if kwargs.get('include_images', True) and self.pdf_converter.images:
                # Determine if we should embed or save images separately
                embed_images = kwargs.get('embed_images', False)
                
                if embed_images:
                    # Embed images as base64
                    html += '<section id="images"><h2>Images</h2>\n'
                    
                    for i, img_data in enumerate(self.pdf_converter.images):
                        if 'image' in img_data:
                            import base64
                            b64_img = base64.b64encode(img_data['image']).decode('utf-8')
                            mime_type = f"image/{img_data['ext']}"
                            
                            html += f'<figure>\n'
                            html += f'<img src="data:{mime_type};base64,{b64_img}" '
                            html += f'alt="Image {i+1}" />\n'
                            html += f'<figcaption>Image {i+1}</figcaption>\n'
                            html += '</figure>\n'
                            
                    html += '</section>\n'
                else:
                    # Save images to an 'images' directory
                    images_dir = os.path.splitext(output_path)[0] + '_images'
                    os.makedirs(images_dir, exist_ok=True)
                    
                    html += '<section id="images"><h2>Images</h2>\n'
                    
                    for i, img_data in enumerate(self.pdf_converter.images):
                        if 'image' in img_data:
                            img_filename = f'image_{i}.{img_data["ext"]}'
                            img_path = os.path.join(images_dir, img_filename)
                            
                            # Save image to file
                            with open(img_path, 'wb') as img_file:
                                img_file.write(img_data['image'])
                            
                            rel_path = f'{os.path.basename(images_dir)}/{img_filename}'
                            
                            html += f'<figure>\n'
                            html += f'<img src="{rel_path}" alt="Image {i+1}" />\n'
                            html += f'<figcaption>Image {i+1}</figcaption>\n'
                            html += '</figure>\n'
                            
                    html += '</section>\n'
            
            # Close the HTML structure
            html += '''
                </main>
                <footer>
                    <p>Generated by Book Converter</p>
                </footer>
            </body>
            </html>
            '''
            
            # Prettify HTML
            soup = BeautifulSoup(html, 'html.parser')
            pretty_html = soup.prettify()
            
            # Write to file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(pretty_html)
            
            logger.info(f"Successfully converted to HTML: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error converting to HTML: {e}")
            raise ValueError(f"Failed to convert to HTML: {e}")


class TextConverter(BaseFormatConverter):
    """
    Converter for plain text format.
    """
    
    def convert(self, output_path: str, **kwargs) -> str:
        """
        Convert the PDF to plain text format.
        
        Args:
            output_path (str): Path to save the output text file
            **kwargs: Additional options:
                - include_metadata (bool): Whether to include metadata
                - line_width (int): Maximum line width
                
        Returns:
            str: Path to the output text file
        """
        try:
            # Extract text from the PDF
            text_by_page = self.pdf_converter.extract_text(
                include_tables=kwargs.get('include_tables', True),
                detect_columns=kwargs.get('detect_columns', True)
            )
            
            # Set up output text
            output_text = ""
            
            # Add metadata if requested
            if kwargs.get('include_metadata', True):
                output_text += "=" * 50 + "\n"
                output_text += "DOCUMENT INFORMATION\n"
                output_text += "=" * 50 + "\n\n"
                
                for key, value in self.pdf_converter.metadata.items():
                    if value:  # Only include non-empty metadata
                        output_text += f"{key.capitalize()}: {value}\n"
                        
                output_text += "\n" + "=" * 50 + "\n\n"
            
            # Process table of contents if available
            pdf_toc = self.pdf_converter.toc
            has_toc = pdf_toc and len(pdf_toc) > 0
            
            # Add TOC if available
            if has_toc and kwargs.get('include_toc', True):
                output_text += "TABLE OF CONTENTS\n"
                output_text += "=" * 50 + "\n\n"
                
                for level, title, page in pdf_toc:
                    # Indentation based on level
                    indent = '  ' * (level - 1)
                    output_text += f"{indent}{title} ... {page}\n"
                    
                output_text += "\n" + "=" * 50 + "\n\n"
            
            # Add content based on TOC or page by page
            if has_toc:
                # Use TOC to structure the document
                current_page = 0
                for i, (level, title, page) in enumerate(pdf_toc):
                    # Add section heading
                    output_text += "\n" + "=" * 50 + "\n"
                    output_text += f"{title}\n"
                    output_text += "=" * 50 + "\n\n"
                    
                    # Get content for this section
                    next_page = pdf_toc[i+1][2] if i+1 < len(pdf_toc) else len(text_by_page)
                    section_content = "".join(text_by_page[page-1:next_page-1])
                    
                    # Add to output
                    output_text += section_content + "\n\n"
            else:
                # No TOC, just add content page by page
                for i, page_text in enumerate(text_by_page):
                    output_text += page_text
                    
                    # Add page break between pages
                    if i < len(text_by_page) - 1:
                        output_text += "\n\n" + "-" * 50 + "\n\n"
            
            # Wrap lines if requested
            line_width = kwargs.get('line_width')
            if line_width:
                wrapped_lines = []
                for line in output_text.splitlines():
                    if len(line) > line_width:
                        # Simple line wrapping algorithm
                        current_line = ""
                        for word in line.split():
                            if len(current_line) + len(word) + 1 <= line_width:
                                current_line += (" " + word if current_line else word)
                            else:
                                wrapped_lines.append(current_line)
                                current_line = word
                        if current_line:
                            wrapped_lines.append(current_line)
                    else:
                        wrapped_lines.append(line)
                        
                output_text = "\n".join(wrapped_lines)
            
            # Write to file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(output_text)
            
            logger.info(f"Successfully converted to TXT: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error converting to TXT: {e}")
            raise ValueError(f"Failed to convert to TXT: {e}")


class MarkdownConverter(BaseFormatConverter):
    """
    Converter for Markdown format.
    """
    
    def convert(self, output_path: str, **kwargs) -> str:
        """
        Convert the PDF to Markdown format.
        
        Args:
            output_path (str): Path to save the output Markdown file
            **kwargs: Additional options:
                - include_metadata (bool): Whether to include metadata
                - include_toc (bool): Whether to include table of contents
                - include_images (bool): Whether to include images
                
        Returns:
            str: Path to the output Markdown file
        """
        try:
            # Extract text from the PDF
            text_by_page = self.pdf_converter.extract_text(
                include_tables=kwargs.get('include_tables', True),
                detect_columns=kwargs.get('detect_columns', True)
            )
            
            # Set up output text
            md_text = ""
            
            # Get metadata
            title = self.pdf_converter.metadata.get('title', 'Untitled Document')
            author = self.pdf_converter.metadata.get('author', 'Unknown Author')
            
            # Add title and metadata
            md_text += f"# {title}\n\n"
            
            if kwargs.get('include_metadata', True):
                md_text += "_Document Information_\n\n"
                md_text += "| Field | Value |\n"
                md_text += "|-------|-------|\n"
                
                for key, value in self.pdf_converter.metadata.items():
                    if value:  # Only include non-empty metadata
                        md_text += f"| {key.capitalize()} | {value} |\n"
                        
                md_text += "\n---\n\n"
            
            # Process table of contents if available
            pdf_toc = self.pdf_converter.toc
            has_toc = pdf_toc and len(pdf_toc) > 0
            
            # Add TOC if available
            if has_toc and kwargs.get('include_toc', True):
                md_text += "## Table of Contents\n\n"
                
                for level, title, page in pdf_toc:
                    # Indentation based on level
                    indent = '  ' * (level - 1)
                    md_text += f"{indent}- [{title}](#page-{page})\n"
                    
                md_text += "\n---\n\n"
            
            # Add content based on TOC or page by page
            if has_toc:
                # Use TOC to structure the document
                current_page = 0
                for i, (level, title, page) in enumerate(pdf_toc):
                    # Add section heading
                    md_text += f"\n{'#' * (level + 1)} {title} {{#page-{page}}}\n\n"
                    
                    # Get content for this section
                    next_page = pdf_toc[i+1][2] if i+1 < len(pdf_toc) else len(text_by_page)
                    section_content = "".join(text_by_page[page-1:next_page-1])
                    
                    # Convert to paragraphs
                    paragraphs = section_content.split('\n\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            md_text += f"{para_text.strip()}\n\n"
            else:
                # No TOC, just add content page by page
                for i, page_text in enumerate(text_by_page):
                    if i == 0:
                        md_text += f"## Content {{#page-{i+1}}}\n\n"
                    else:
                        md_text += f"## Page {i+1} {{#page-{i+1}}}\n\n"
                    
                    # Convert to paragraphs
                    paragraphs = page_text.split('\n\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            md_text += f"{para_text.strip()}\n\n"
            
            # Handle images
            if kwargs.get('include_images', True) and self.pdf_converter.images:
                # Save images to a directory
                images_dir = os.path.splitext(output_path)[0] + '_images'
                os.makedirs(images_dir, exist_ok=True)
                
                md_text += "## Images\n\n"
                
                for i, img_data in enumerate(self.pdf_converter.images):
                    if 'image' in img_data:
                        img_filename = f'image_{i}.{img_data["ext"]}'
                        img_path = os.path.join(images_dir, img_filename)
                        
                        # Save image to file
                        with open(img_path, 'wb') as img_file:
                            img_file.write(img_data['image'])
                        
                        rel_path = f'{os.path.basename(images_dir)}/{img_filename}'
                        
                        md_text += f"![Image {i+1}]({rel_path})\n\n"
                        md_text += f"*Image {i+1}*\n\n"
            
            # Write to file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(md_text)
            
            logger.info(f"Successfully converted to Markdown: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error converting to Markdown: {e}")
            raise ValueError(f"Failed to convert to Markdown: {e}")


class MOBIConverter(BaseFormatConverter):
    """
    Converter for MOBI format.
    """
    
    def convert(self, output_path: str, **kwargs) -> str:
        """
        Convert the PDF to MOBI format.
        This requires Calibre's ebook-convert tool to be installed on the system.
        
        Args:
            output_path (str): Path to save the output MOBI file
            **kwargs: Additional options from EPUBConverter
                
        Returns:
            str: Path to the output MOBI file
            
        Raises:
            ValueError: If conversion fails or ebook-convert is not installed
        """
        try:
            # First convert to EPUB
            with tempfile.TemporaryDirectory() as temp_dir:
                epub_path = os.path.join(temp_dir, 'temp.epub')
                
                # Use the EPUB converter
                epub_converter = EPUBConverter(self.pdf_converter)
                epub_result = epub_converter.convert(epub_path, **kwargs)
                
                logger.info(f"EPUB created, now attempting conversion to MOBI")
                
                # Check if Calibre's ebook-convert is available
                if self._check_calibre_available():
                    # Convert using Calibre
                    self._convert_with_calibre(epub_path, output_path)
                else:
                    # Fallback: just copy the EPUB file
                    logger.warning("Calibre ebook-convert not found; MOBI conversion unavailable")
                    logger.warning("Copying EPUB as a fallback (rename to .mobi extension)")
                    shutil.copy2(epub_path, output_path)
                    logger.info("Note: To enable proper MOBI conversion, install Calibre and ensure 'ebook-convert' is in your PATH")
                
                if os.path.exists(output_path):
                    logger.info(f"Successfully created MOBI (or fallback): {output_path}")
                    return output_path
                else:
                    raise ValueError("Failed to create output file")
                
        except Exception as e:
            logger.error(f"Error converting to MOBI: {e}")
            raise ValueError(f"Failed to convert to MOBI: {e}")

    def _check_calibre_available(self) -> bool:
        """
        Check if Calibre's ebook-convert is available in the system.
        
        Returns:
            bool: True if available, False otherwise
        """
        try:
            # Try to run ebook-convert with --version
            result = subprocess.run(
                ['ebook-convert', '--version'], 
                stdout=subprocess.PIPE, 
                stderr=subprocess.PIPE,
                timeout=5
            )
            return result.returncode == 0
        except (subprocess.SubprocessError, FileNotFoundError):
            return False
    
    def _convert_with_calibre(self, epub_path: str, mobi_path: str) -> None:
        """
        Convert EPUB to MOBI using Calibre's ebook-convert tool.
        
        Args:
            epub_path (str): Path to input EPUB file
            mobi_path (str): Path to output MOBI file
            
        Raises:
            ValueError: If conversion fails
        """
        try:
            logger.info("Converting EPUB to MOBI using Calibre's ebook-convert...")
            
            # Run the conversion
            result = subprocess.run(
                ['ebook-convert', epub_path, mobi_path],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                timeout=60  # 1 minute timeout
            )
            
            # Check if successful
            if result.returncode != 0:
                logger.error(f"Calibre conversion failed: {result.stderr}")
                raise ValueError(f"ebook-convert failed: {result.stderr}")
                
            logger.info("Calibre conversion successful")
            
        except subprocess.TimeoutExpired:
            logger.error("Calibre conversion timed out")
            raise ValueError("MOBI conversion timed out")
        except subprocess.SubprocessError as e:
            logger.error(f"Calibre process error: {e}")
            raise ValueError(f"MOBI conversion failed: {e}")
